from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph


SOURCE = Path("/Users/monkeyyey/Downloads/Final_Report.docx")
TARGET = Path("/Users/monkeyyey/Documents/Code/log_pipeline_updated/Final_Report - merged analyzer.docx")


def replace_text(paragraph, text: str) -> None:
    for run in paragraph.runs:
        run.text = ""
    if paragraph.runs:
        paragraph.runs[0].text = text
    else:
        paragraph.add_run(text)


def insert_paragraph_after(paragraph, text: str, style: str | None = None):
    new_p = OxmlElement("w:p")
    paragraph._element.addnext(new_p)
    inserted = Paragraph(new_p, paragraph._parent)
    replace_text(inserted, text)
    if style:
        inserted.style = style
    return inserted


def insert_after_by_text(doc: Document, anchor_text: str, entries: list[tuple[str, str | None]]):
    for paragraph in doc.paragraphs:
        if paragraph.text.strip() == anchor_text:
            current = paragraph
            for text, style in entries:
                current = insert_paragraph_after(current, text, style)
            return
    raise ValueError(f"Anchor not found: {anchor_text}")


def find_paragraph(doc: Document, startswith: str):
    for paragraph in doc.paragraphs:
        if paragraph.text.strip().startswith(startswith):
            return paragraph
    raise ValueError(f"Paragraph not found: {startswith}")


def insert_after_by_startswith(doc: Document, startswith: str, entries: list[tuple[str, str | None]]):
    paragraph = find_paragraph(doc, startswith)
    current = paragraph
    for text, style in entries:
        current = insert_paragraph_after(current, text, style)


def main() -> None:
    doc = Document(str(SOURCE))

    replace_text(
        find_paragraph(doc, "Machine-generated logs are produced"),
        (
            "Machine-generated logs are produced in a wide variety of formats across different systems, "
            "services, vendors, and semiconductor tools. Inconsistent structure, mixed encodings, varying "
            "timestamp conventions, and malformed payloads make direct downstream analysis difficult. This "
            "project addresses the problem with a two-part Flask application. The first part is a multi-format "
            "ETL pipeline that detects log structure, parses eleven formats into a canonical LogRecord schema, "
            "normalises core fields, flags corruption, and stages cleaned outputs. The second part is a SQLite "
            "Analyzer dashboard that supports interactive inspection of mapping registries and processed outputs. "
            "After merging, both the analyzer and pipeline dashboard run under one Flask server at /analyzer/ "
            "and /pipeline/, linked through shared navigation. Together they support both automated preparation "
            "and analyst-facing validation of semiconductor machinery log data."
        ),
    )

    replace_text(
        find_paragraph(doc, "The pipeline is designed to satisfy the following objectives:"),
        (
            "The combined application is designed to satisfy the following objectives:"
        ),
    )

    objective_anchor = find_paragraph(doc, "Provide a browser-based dashboard for interactive exploration")
    replace_text(
        objective_anchor,
        "Provide a browser-based dashboard for interactive exploration, SQLite inspection, and mapping quality review without requiring separate database infrastructure.",
    )
    insert_paragraph_after(
        objective_anchor,
        "Expose both the analyzer and the pipeline dashboard through one Flask deployment so operators can move between ingestion and inspection without changing applications.",
        objective_anchor.style.name,
    )

    replace_text(
        find_paragraph(doc, "The pipeline is structured as a layered ETL system."),
        (
            "The project is structured as a layered ETL system with a merged Flask presentation layer. "
            "Raw log data enters through a single service facade (LogIngestionService), flows into a staging "
            "area (StagingArea), and can then be inspected through a SQLite-backed analyzer. The combined web "
            "application exposes two mounted routes: /pipeline/ for ingestion and parsing workflows, and "
            "/analyzer/ for browsing SQLite outputs, mapping registries, and quality metrics."
        ),
    )

    replace_text(
        find_paragraph(doc, "The Flask application (app.py) wraps the pipeline in a REST API."),
        (
            "The web layer now consists of a top-level Flask composition app that mounts two sub-applications. "
            "The /pipeline/ route serves the original ingestion dashboard, which wraps the parsing pipeline in a "
            "REST API for file upload, format detection, processing, staging statistics, and export. The "
            "/analyzer/ route serves a SQLite inspection dashboard oriented around semiconductor machinery-log "
            "mapping registries and processed outputs. The analyzer supports SQLite upload, database switching, "
            "database removal, table browsing, row search, schema profiling, chart generation, signal coverage, "
            "confidence analysis, format distribution, recent signature activity, and a watchlist of mappings "
            "that need attention. Both pages share a dark visual theme and a navigation bar so users can move "
            "between ingestion and inspection within one Flask server."
        ),
    )

    insert_after_by_startswith(
        doc,
        "The web layer now consists of a top-level Flask composition app",
        [
            ("4.7 SQLite Analyzer Dashboard", "Heading 3"),
            (
                "The analyzer is intended for registry-style SQLite datasets generated by the pipeline or by "
                "downstream semantic processing. It is especially useful when a table contains fields such as "
                "signature, format_type, regex_patterns, confidence, hit_count, created_at, and updated_at. "
                "When these fields are present, the analyzer computes average confidence, observed hits, format "
                "mix, tracked signal coverage, recent signature edits, and heuristic watchlist flags.",
                "normal",
            ),
            (
                "Signal coverage is computed across eight tracked semiconductor telemetry fields: timestamp, "
                "tool_id, alarm_code, event_code, lot_id, recipe, chamber, and wafer_id. Coverage is a "
                "breadth-of-representation metric across the selected table, not a claim that every mapping "
                "captures every field. For example, 8/8 coverage means all eight signal categories appear "
                "somewhere in the dataset.",
                "normal",
            ),
            (
                "Mappings are flagged for attention when confidence is below 0.75, when hit_count is high "
                "but confidence is still below 0.90, or when regex_patterns extracts only zero or one fields. "
                "These heuristics are designed to surface brittle mappings that may hide alarm context, miss "
                "tool identity, or weaken downstream analytics.",
                "normal",
            ),
            ("4.8 Merged Route Structure", "Heading 3"),
            (
                "The final deployment uses a top-level Flask application with DispatcherMiddleware. The root "
                "path redirects to /analyzer/. The analyzer is mounted at /analyzer/, and the imported "
                "pipeline dashboard is mounted at /pipeline/. This preserves separation of concerns in code "
                "while presenting one coherent application to the user.",
                "normal",
            ),
        ],
    )

    replace_text(
        find_paragraph(doc, "All REST endpoints return JSON."),
        (
            "All REST endpoints return JSON. Error responses follow {\"error\": \"<message>\"} with an "
            "appropriate HTTP status code. After merging, the route surface is split between /pipeline/api/... "
            "for ingestion operations and /analyzer/api/... for SQLite inspection and analyzer metrics."
        ),
    )

    insert_after_by_startswith(
        doc,
        "DELETE /api/staging",
        [
            ("Analyzer endpoints:", "normal"),
            (
                "GET /analyzer/api/databases                    POST /analyzer/api/databases/upload  "
                "(multipart/form-data, field: \"file\")                    DELETE /analyzer/api/databases/<db_id>",
                "normal",
            ),
            (
                "GET /analyzer/api/tables?db=<db_id>                    GET /analyzer/api/schema/<table>?db=<db_id>                    "
                "GET /analyzer/api/rows/<table>?db=<db_id>&limit=50&search=",
                "normal",
            ),
            (
                "GET /analyzer/api/profile/<table>?db=<db_id>                    "
                "GET /analyzer/api/chart/<table>?db=<db_id>&dimension=<col>&metric=<col>&agg=count                    "
                "GET /analyzer/api/analyzer/<table>?db=<db_id>",
                "normal",
            ),
        ],
    )

    replace_text(
        find_paragraph(doc, "State persistence: The Flask application maintains pipeline state"),
        (
            "State persistence: The merged Flask application mixes two persistence models. The pipeline "
            "dashboard maintains staging state in process-local Python objects and therefore loses that state "
            "on server restart. The analyzer persists uploaded SQLite files in uploaded_dbs/ and can reopen "
            "them across sessions, but it does not itself persist browser selections or dashboard state."
        ),
    )

    replace_text(
        find_paragraph(doc, "Concurrency: The development server"),
        (
            "Concurrency: The development server (flask run / app.run(debug=True)) is not intended for "
            "multi-user production use. The combined application mounts two sub-apps in one process, so "
            "concurrent requests across analyzer and pipeline views should be served behind a production WSGI "
            "server such as Gunicorn or Waitress if shared access is required."
        ),
    )

    replace_text(
        find_paragraph(doc, "Upload size: The MAX_CONTENT_LENGTH"),
        (
            "Upload size: Request size limits now depend on the mounted application. The pipeline retains its "
            "50 MB upload guard for raw log processing, while the analyzer accepts SQLite uploads subject to "
            "filesystem capacity and any Flask-level request limits configured for deployment."
        ),
    )

    replace_text(
        find_paragraph(doc, "Security: The application is intended"),
        (
            "Security: The merged application is intended for local development and competition demonstration. "
            "Uploaded filenames are sanitised, but no authentication, role separation, rate limiting, or "
            "content sandboxing is applied. Because the analyzer opens user-supplied SQLite files and the "
            "pipeline accepts arbitrary log uploads, the application should not be exposed to untrusted networks."
        ),
    )

    insert_after_by_startswith(
        doc,
        "The full system is structured as two independent layers.",
        [
            (
                "The analyzer sits above these layers as an inspection surface rather than a transformation "
                "stage. It reads the SQLite outputs, registry tables, or intermediate databases produced by "
                "earlier stages and gives analysts a fast way to validate mapping quality, signal coverage, "
                "and recent changes before those datasets are consumed downstream.",
                "normal",
            )
        ],
    )

    insert_after_by_startswith(
        doc,
        "6. ML-ready feature extraction:",
        [
            (
                "7. Analyzer-driven governance: The analyzer could be extended with approval workflows, owner "
                "assignment, regression baselines, and historical drift tracking so that mapping quality is "
                "treated as a first-class operational concern rather than an ad hoc manual review task.",
                "normal",
            )
        ],
    )

    replace_text(
        find_paragraph(doc, "The two-layer architecture reflects a deliberate separation of concerns."),
        (
            "The final system reflects a deliberate separation of concerns. The Format Parsing Layer handles "
            "the syntax problem of turning heterogeneous raw logs into consistent Python records. The Semantic "
            "Pipeline handles the semantic problem of resolving field meaning, normalising units, and scoring "
            "equipment behaviour. The merged Flask application adds an analyst-facing inspection layer on top "
            "of both, allowing users to move between ingestion and SQLite analysis without changing tools."
        ),
    )

    replace_text(
        find_paragraph(doc, "The system as a whole addresses the full data preparation workflow"),
        (
            "The system as a whole addresses the full data preparation workflow for semiconductor equipment "
            "logs: from raw, heterogeneous log files in eleven formats to a normalised, deduplicated, "
            "anomaly-scored, queryable SQLite database, and finally to an interactive analyzer that exposes "
            "mapping confidence, signal coverage, recent signature edits, and tables that need attention. "
            "This merged design improves both automation and review by keeping preparation, inspection, and "
            "iterative tuning in one coherent application."
        ),
    )

    insert_after_by_text(
        doc,
        "Appendix C: Observed Pipeline Performance",
        [
            (
                "Note: the analyzer dashboard was added after the original pipeline benchmark runs were "
                "collected. The metrics in Appendix C therefore measure ingestion, parsing, staging, and "
                "semantic processing rather than browser-side inspection latency.",
                "normal",
            )
        ],
    )

    TARGET.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(TARGET))
    print(TARGET)


if __name__ == "__main__":
    main()
